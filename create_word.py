import os
try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call(["python", "-m", "pip", "install", "python-docx"])
    from docx import Document

text = """BÁO CÁO TÍNH NĂNG & ĐIỂM NỔI BẬT CỦA DỰ ÁN GIA SƯ ẢO

Đây là điểm ăn tiền nhất của Gia Sư Ảo để bạn tự tin nhấn mạnh trong buổi báo cáo. Việc so sánh này sẽ làm bật lên giá trị của dự án so với việc "chỉ dùng ChatGPT là xong".

Phần 1: Sự Cá nhân hóa (Personalization) nằm ở đâu?

- Cá nhân hóa theo Cấp học & Sách Giáo Khoa (Context-Aware): Khi bạn dùng ChatGPT, nó không biết bạn là ai. Nhưng với Gia Sư Ảo, hệ thống biết chính xác học sinh đang học Lớp mấy (VD: Lớp 4), Môn gì (VD: Khoa học) và Sách gì (VD: Kết nối tri thức).
- Cá nhân hóa Văn phong & Xưng hô: Với học sinh Tiểu học, AI xưng "Thầy/Cô - Em", dùng ngôn từ cổ vũ dễ thương. Nhưng với học sinh THPT, AI tự động chuyển sang văn phong mạch lạc, học thuật và chuyên sâu hơn.
- Cá nhân hóa Giao diện (Dynamic UI): Giao diện và màu sắc app tự động biến đổi theo độ tuổi. Học sinh lớp 3 chỉ thấy các môn như Tự nhiên Xã hội, trong khi học sinh lớp 12 sẽ thấy các môn chuyên sâu Vật Lý, Hóa Học.
- Gợi ý Học tập May đo (Tailored Suggestions): Các nút gợi ý cuối câu trả lời không phải là random. Nó được AI tự động phân tích từ chính điểm yếu trong câu hỏi vừa xong của học sinh để đưa ra gợi ý tiếp theo phù hợp nhất.

Phần 2: Điểm Vượt Trội Hơn ChatGPT và Các AI/App Khác (USPs)

Thứ nhất: Có Ranh giới Giáo dục (Guardrails) - Không Dạy Hư Học Sinh
- ChatGPT/Các AI khác: Học sinh bảo "Viết cho tôi bài văn tả con mèo" hoặc "Giải bài này cho tôi", AI sẽ làm hộ từ A-Z -> Học sinh chép bài thụ động.
- Gia Sư Ảo: Được thiết lập ranh giới sư phạm. Nó sẽ từ chối giải bài hộ hoặc từ chối giải kiến thức vượt cấp (Lớp 1 hỏi toán lớp 12). Thay vào đó, nó đóng vai trò Người dẫn dắt, đưa ra dàn ý hoặc công thức để học sinh tự tư duy tìm ra đáp án.

Thứ hai: Tương tác Đa giác quan (Khép kín)
- Các App khác: Chỉ thuần Text (Văn bản).
- Gia Sư Ảo: Học sinh có thể Chụp ảnh bài tập gửi lên (Mắt) -> AI đọc hiểu và phân tích đề -> AI giảng bài bằng Giọng nói tiếng Việt/Tiếng Anh bản xứ (Tai) -> Trình bày công thức Toán học bằng chuẩn KaTeX cực đẹp mắt. Một trải nghiệm trọn vẹn không cần dùng đến 3-4 app khác nhau.

Thứ ba: Dữ liệu Đáng Tin Cậy (Chống Ảo giác AI)
- ChatGPT: Thường xuyên "bịa" ra thông tin (Hallucination) hoặc lấy thông tin trôi nổi trên mạng, rất nguy hiểm cho học sinh.
- Gia Sư Ảo: Sử dụng công nghệ RAG (Retrieval-Augmented Generation) để neo dữ liệu vào chính sách giáo khoa và tài liệu của Bộ GD&ĐT do nhà trường/admin cung cấp. AI chỉ được phép giảng dạy dựa trên nguồn tài liệu chuẩn mực này.

Thứ tư: Gamification (Biến việc học AI thành Trò chơi)
- Dùng các AI khác rất nhàm chán và cô đơn. Nhưng Gia Sư Ảo kết nối việc học AI với Hệ thống Điểm thưởng (SP), Chuỗi ngày học liên tục (Streak) và Đồng hồ bấm giờ tập trung (Pomodoro). Mỗi phút trò chuyện với AI đều được ghi nhận lại, tạo động lực to lớn cho học sinh.

Tóm lại: Gia Sư Ảo không phải là một công cụ để hỏi đáp thụ động như ChatGPT. Nó là một Môi trường sư phạm thông minh, biết học sinh là ai, bảo vệ học sinh khỏi thông tin độc hại, ép học sinh phải tư duy và thưởng cho sự nỗ lực của họ.

-----------------------------------------------------------
TÓM TẮT NGẮN GỌN DÀNH CHO BÁO CÁO

1. Trí tuệ nhân tạo (AI Tutor Core)
- Học theo bối cảnh: Tự động điều chỉnh giọng văn và kiến thức khớp với Lớp và Môn học của từng học sinh. Từ chối trả lời nếu hỏi vượt cấp.
- Gợi ý thông minh (Chips): Tự động sinh ra 2-3 câu hỏi gợi ý sau mỗi lần AI trả lời để học sinh bấm hỏi tiếp.
- Xử lý Toán học: Hiển thị công thức Toán, Lý, Hóa chuẩn quốc tế cực đẹp bằng KaTeX.

2. Đa phương tiện & Xử lý tài liệu
- Đọc hiểu File (RAG): Cho phép upload ảnh hoặc file PDF bài tập, AI sẽ tự động đọc, tóm tắt và giải chi tiết.
- Đọc bằng giọng nói (TTS): Tự động nhận diện ngôn ngữ. Tiếng Việt đọc bằng giọng chuẩn FPT/Zalo, Tiếng Anh đọc bằng giọng bản xứ Google TTS.

3. Động lực học tập (Gamification)
- Dữ liệu Real-time: Liên kết cơ sở dữ liệu thật (Supabase) để theo dõi thời gian học, tính điểm thưởng (SP) và chuỗi ngày học liên tục (Streak).

4. Kỹ thuật & Hiệu năng
- Tốc độ: Tích hợp Streaming (AI gõ chữ theo thời gian thực) không có độ trễ.
- Dùng thử (Guest Mode): Cho phép trải nghiệm nhanh 3 phút không cần đăng ký.
- Kiến trúc Đám mây: Nâng cấp thành công Backend lên Vercel Serverless, chạy ổn định 24/7 và đồng bộ API chặt chẽ. Đã dọn sạch 100% rác trong mã nguồn.

1. Sự "Cá nhân hóa" (Personalization) nằm ở đâu?
Cá nhân hóa theo Cấp học & Sách Giáo Khoa (Context-Aware): Khi bạn dùng ChatGPT, nó không biết bạn là ai. Nhưng với "Gia Sư Ảo", hệ thống biết chính xác học sinh đang học Lớp mấy (VD: Lớp 4), Môn gì (VD: Khoa học) và Sách gì (VD: Kết nối tri thức).
Cá nhân hóa Văn phong & Xưng hô: Với học sinh Tiểu học, AI xưng "Thầy/Cô - Em", dùng ngôn từ cổ vũ dễ thương kèm Emoji (🌟, 👏). Nhưng với học sinh THPT, AI tự động chuyển sang văn phong mạch lạc, học thuật và chuyên sâu hơn.
Cá nhân hóa Giao diện (Dynamic UI): Giao diện và màu sắc app tự động biến đổi theo độ tuổi. Học sinh lớp 3 chỉ thấy các môn như "Tự nhiên Xã hội", trong khi học sinh lớp 12 sẽ thấy các môn chuyên sâu "Vật Lý, Hóa Học".
Gợi ý Học tập May đo (Tailored Suggestions): Các nút gợi ý cuối câu trả lời không phải là random. Nó được AI tự động phân tích từ chính điểm yếu trong câu hỏi vừa xong của học sinh để đưa ra gợi ý tiếp theo phù hợp nhất.
2. Điểm Vượt Trội Hơn ChatGPT và Các AI/App Khác (USPs)
Thứ nhất: Có "Ranh giới Giáo dục" (Guardrails) - Không Dạy Hư Học Sinh

ChatGPT/Các AI khác: Học sinh bảo "Viết cho tôi bài văn tả con mèo" hoặc "Giải bài này cho tôi", AI sẽ làm hộ từ A-Z -> Học sinh chép bài thụ động.
Gia Sư Ảo: Được thiết lập "Ranh giới sư phạm". Nó sẽ từ chối giải bài hộ hoặc từ chối giải kiến thức vượt cấp (Lớp 1 hỏi toán lớp 12). Thay vào đó, nó đóng vai trò "Người dẫn dắt", đưa ra dàn ý hoặc công thức để học sinh tự tư duy tìm ra đáp án.
Thứ hai: Tương tác Đa giác quan (Khép kín)

Các App khác: Chỉ thuần Text (Văn bản).
Gia Sư Ảo: Học sinh có thể Chụp ảnh bài tập gửi lên (Mắt) → AI đọc hiểu và phân tích đề → AI giảng bài bằng Giọng nói tiếng Việt/Tiếng Anh bản xứ (Tai) → Trình bày công thức Toán học bằng chuẩn KaTeX cực đẹp mắt. Một trải nghiệm trọn vẹn không cần dùng đến 3-4 app khác nhau.
Thứ ba: Dữ liệu Đáng Tin Cậy (Chống "Ảo giác AI")

ChatGPT: Thường xuyên "bịa" ra thông tin (Hallucination) hoặc lấy thông tin trôi nổi trên mạng, rất nguy hiểm cho học sinh.
Gia Sư Ảo: Sử dụng công nghệ RAG (Retrieval-Augmented Generation) để neo dữ liệu vào chính sách giáo khoa và tài liệu của Bộ GD&ĐT do nhà trường/admin cung cấp. AI chỉ được phép giảng dạy dựa trên nguồn tài liệu chuẩn mực này.
Thứ tư: Gamification (Biến việc học AI thành Trò chơi)

Dùng các AI khác rất nhàm chán và cô đơn. Nhưng "Gia Sư Ảo" kết nối việc học AI với Hệ thống Điểm thưởng (SP), Chuỗi ngày học liên tục (Streak) và Đồng hồ bấm giờ tập trung (Pomodoro). Mỗi phút trò chuyện với AI đều được ghi nhận lại, tạo động lực to lớn cho học sinh.
👉 Tóm lại: "Gia Sư Ảo" không phải là một công cụ để hỏi đáp thụ động như ChatGPT. Nó là một Môi trường sư phạm thông minh, biết học sinh là ai, bảo vệ học sinh khỏi thông tin độc hại, ép học sinh phải tư duy và thưởng cho sự nỗ lực của họ.

. Mô hình nhúng (Embedding Model)
Sử dụng mô hình models/text-embedding-004 (hoặc fallback về models/embedding-001) của Google Gemini.
Nó có nhiệm vụ "số hóa" (vector hóa) các tài liệu sách giáo khoa và các file bài tập mà học sinh tải lên thành các dãy số nhiều chiều để máy tính hiểu được ý nghĩa ngữ nghĩa (Semantic Meaning).
2. Cơ sở dữ liệu Vector (Vector Database)
Hệ thống sử dụng Supabase PostgreSQL kết hợp với phần mở rộng pgvector.
Khi một file PDF (như sách giáo khoa hoặc đề bài) được upload, Backend sẽ chẻ nhỏ văn bản (Chunking), mã hóa thành các vector và lưu trữ vào bảng documents trên Supabase.
3. Thuật toán Truy xuất (Retrieval)
Sử dụng Vector Similarity Search (Tìm kiếm độ tương đồng Vector).
Khi học sinh hỏi một câu (VD: "Động lượng là gì?"), câu hỏi cũng được chuyển thành vector. Sau đó, Supabase sử dụng hàm RPC (Remote Procedure Call) match_documents để quét toàn bộ CSDL và lôi ra những đoạn văn bản/công thức có độ tương đồng ngữ nghĩa cao nhất với câu hỏi.
4. Mô hình Sinh văn bản (Generator LLM)
Sử dụng mô hình Google Gemini 2.5 Flash (có fallback về 1.5 Pro).
Thay vì để AI tự do trả lời theo trí nhớ mạng Internet dễ bị sai lệch (Ảo giác AI - Hallucination), Backend sẽ "bơm" trực tiếp các đoạn tài liệu vừa trích xuất được ở bước 3 vào Prompt hệ thống.
Lệnh Prompt RAG ép buộc: "Sử dụng các tài liệu sau để trả lời. Nếu không có trong tài liệu, hãy dựa vào kiến thức nền tảng nhưng phải tuân thủ chuẩn SGK cấp học hiện tại".
👉 Tóm lại: Giải pháp RAG này là sự kết hợp hoàn hảo giữa Gemini Embedding (Mã hóa) + Supabase pgvector (Lưu trữ/Tìm kiếm) + Gemini 2.5 Flash (Sinh câu trả lời), tạo ra một trợ lý AI thông minh nhưng cực kỳ bám sát sách giáo khoa!

Đọc PDF/Ảnh: Dùng markitdown của Microsoft và pypdf cực kỳ chính xác.
Cắt văn bản (Chunking): Tự viết thuật toán cắt text.
Mã hóa (Embedding): Gọi trực tiếp genai.embed_content của Google.
Tìm kiếm (Vector Search): Tự viết hàm SQL (RPC match_documents) trên Supabase pgvector.
"""

doc = Document()
doc.add_heading('BÁO CÁO DỰ ÁN GIA SƯ ẢO', 0)

for paragraph in text.split('\n'):
    if paragraph.strip() == '':
        continue
    if paragraph.startswith('Phần') or paragraph.startswith('TÓM TẮT') or paragraph.startswith('Thứ') or paragraph[0].isdigit():
        doc.add_heading(paragraph, level=2)
    else:
        doc.add_paragraph(paragraph)

output_path = r"c:\Users\Admin\Downloads\BaoCao_GiaSuAo.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
